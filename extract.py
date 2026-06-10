# -*- coding: utf-8 -*-
"""Extract full chapter content (headings + body + captions + IMAGES) and
Seoul-experience cases from the book docx drafts. Images are saved under
static/img/<order>/ and referenced in data.json content blocks."""
import docx, glob, os, unicodedata, json, re, io
from collections import Counter
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from PIL import Image

def norm(s): return unicodedata.normalize('NFC', s or '').strip()

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, 'static', 'img')
SRC = '/Users/jychoi/Library/CloudStorage/OneDrive-개인/업무/02.연구원/2.연구/2026년/2026 정책아카이브/단행본/초본_검수용'

V_IMAGEDATA = '{urn:schemas-microsoft-com:vml}imagedata'
R_ID = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'

META = [
    ('소개말',        0, '여는 글 · Opening',        '소개말: 도시계획이란 무엇인가',        'Prologue: What Is Urban Planning',        '서장'),
    ('20세기 중반',   0, '여는 글 · Opening',        '20세기 중반 서울의 이슈와 도시계획의 방향', 'Mid-20th-Century Seoul: Issues and Directions', '1장'),
    ('밀도체계',      1, 'Part 1. 공간의 뼈대',      '밀도 체계',                          'The Density System',                      '2장'),
    ('중심체계',      1, 'Part 1. 공간의 뼈대',      '중심 체계',                          'The Center System',                       '3장'),
    ('가로체계',      1, 'Part 1. 공간의 뼈대',      '가로 체계',                          'The Street System',                       '4장'),
    ('공공교통',      1, 'Part 1. 공간의 뼈대',      '공공교통 체계',                       'The Public Transit System',               '5장'),
    ('공공공간',      1, 'Part 1. 공간의 뼈대',      '공공공간 체계',                       'The Public Space System',                 '6장'),
    ('자연공간',      1, 'Part 1. 공간의 뼈대',      '자연공간 체계',                       'The Natural Space System',                '7장'),
    ('9장',           2, 'Part 2. 도시의 대사(代謝)', '물 체계',                           'The Water System',                        '8장'),
    ('10장',          2, 'Part 2. 도시의 대사(代謝)', '쓰레기 관리 체계',                    'The Waste Management System',             '9장'),
    ('11장',          2, 'Part 2. 도시의 대사(代謝)', '에너지 관리 체계',                    'The Energy Management System',            '10장'),
    ('12장',          2, 'Part 2. 도시의 대사(代謝)', '홍수·가뭄 관리 체계',                 'The Flood & Drought Management System',   '11장'),
    ('13장',          3, 'Part 3. 삶의 터전',        '경제 체계',                          'The Economic System',                     '12장'),
    ('14장',          3, 'Part 3. 삶의 터전',        '동네 체계',                          'The Neighborhood System',                 '13장'),
    ('15장',          3, 'Part 3. 삶의 터전',        '주거 체계',                          'The Housing System',                      '14장'),
    ('17장',          4, 'Part 4. 계획을 움직이는 힘', '도시계획 거버넌스 체계',              'The Urban Planning Governance System',    '15장'),
    ('18장',          4, 'Part 4. 계획을 움직이는 힘', '도시계획 정보체계와 전자정부',          'Urban Information Systems & e-Government', '16장'),
    ('맺음말',        5, '닫는 글 · Closing',        '맺음말: 성장하는 도시의 동료시민에게',   'Epilogue: To Fellow Citizens of a Growing City', '종장'),
]

CIRCLED = '①②③④⑤⑥⑦⑧⑨⑩'
CORE_RE = re.compile(r'^(핵심\s*(정책|경험)?\s*[①-⑩\d]|Core\s*(Point\s*)?\d|연결\s*\d|Connection\s*\d)')
GENERIC_SUB = re.compile(r'(배경과 맥락|배경|정책 ?내용|실행 ?방식|성과와 한계|성과|한계|시사점|개발도상국)')
CASES_HEAD = re.compile(r'(핵심\s*경험|핵심경험|서울의 경험|서울, 그 증거|핵심 정책)')
NORM_HEAD = re.compile(r'(실무적 원론|원론|국제 표준|계획 및 운영|배분 원칙|정책 제언|패러다임)')
CASE_LABEL = re.compile(r'\(?사례\s*\d')
HANGUL = re.compile(r'[가-힣]')
CAPTION_RE = re.compile(r'^\s*\[?\s*(그림|사진|표|도면|Figure|Fig|Table)\s*[\[\]\d\-]|^\s*(출처|자료|Source)\s*[:：]', re.I)

def split_title(t):
    for sep in ['   |   ', '  |  ', ' | ', '  ｜  ', '|']:
        if sep in t:
            a, b = t.split(sep, 1); return norm(a), norm(b)
    for sep in ['  ·  ', ' · ', ' — ']:
        if sep in t:
            a, b = t.split(sep, 1)
            if b.strip() and not HANGUL.search(b): return norm(a), norm(b)
    return t, ''

def classify(text):
    if CORE_RE.match(text): return 'core'
    if re.match(r'^(도입|들어가며|Introduction|연결)', text): return 'intro'
    if re.match(r'^(규범|The Norm|실무적 원론|원론)', text) or '규범' in text[:6]: return 'norm'
    if re.match(r'^(맺음|맺음말|Closing|Conclusion|결론|책을 닫|책을 마치)', text): return 'closing'
    if re.match(r'^(맥락과 제약|Context|접근과 행동|Approach|결과와 성찰|Result)', text): return 'context'
    if re.match(r'^(참고문헌|References|Reference)', text): return 'refs'
    return 'section'

# ---- per-mode single-paragraph classifiers (return one block dict or None) ----
def cls_heading(p, t, st):
    s = p.style.name
    if 'Heading' in s:
        lvl = int(s.replace('Heading ', '')) if s.replace('Heading ', '').isdigit() else 2
        kr, en = split_title(t)
        return {'t': 'h', 'level': lvl, 'kr': kr, 'en': en, 'kind': classify(t)}
    return {'t': 'p', 'text': t}

NAMED = {'대제목': 1, '1. 대제목': 1, '중제목': 2, '1.1 중제목': 2, '소제목2': 3}
def cls_named(p, t, st):
    s = norm(p.style.name)
    if s in NAMED:
        if s == '중제목' and t.lower().startswith('reference'):
            return {'t': 'h', 'level': 1, 'kr': '참고문헌', 'en': 'References', 'kind': 'refs'}
        lvl = NAMED[s]; kr, en = split_title(t); kind = classify(t)
        if lvl == 1 and re.search(r'핵심\s*정책', t): kind = 'core'
        return {'t': 'h', 'level': lvl, 'kr': kr, 'en': en, 'kind': kind}
    if s in ('표 제목', '그림제목', '그림제목2', '표 내용'):
        return {'t': 'cap', 'text': t}
    if s == 'Reference':
        return {'t': 'ref', 'text': t}
    return {'t': 'p', 'text': t}

def cls_font(p, t, st):
    if p.style.name == 'Caption':
        return {'t': 'cap', 'text': t}
    runs = [r for r in p.runs if r.text.strip()]
    sz = next((r.font.size.pt for r in runs if r.font.size), None)
    is_circ = t[0] in CIRCLED if t else False
    is_major = (sz and sz >= 14) or re.match(r'^\d+장?\.?\s*\S|^\d+\.\d+', t)
    if CASE_LABEL.search(t) and len(t) < 90:
        kr = re.sub(r'^.*?\(?사례\s*\d\)?\s*', '', t).strip() or t
        return {'t': 'h', 'level': 2, 'kr': kr, 'en': '', 'kind': 'core'}
    if is_major and not is_circ:
        kr, en = split_title(t)
        if CASES_HEAD.search(t): st['in_cases'], kind = True, 'casehead'
        elif NORM_HEAD.search(t): st['in_cases'], kind = False, 'norm'
        else: kind = classify(t)
        lvl = 1 if (sz and sz >= 16) or re.match(r'^\d+장|^\d+\.\s', t) else 2
        if (len(kr) < 45 or re.match(r'^\d', kr) or kind != 'section') and not kr.rstrip().endswith(('.', '다', '다.')):
            return {'t': 'h', 'level': lvl, 'kr': kr, 'en': en, 'kind': kind}
        return {'t': 'p', 'text': t}
    if is_circ and len(t) < 90:
        label = t[1:].strip()
        if st['in_cases'] and not GENERIC_SUB.search(label):
            return {'t': 'h', 'level': 2, 'kr': label, 'en': '', 'kind': 'core'}
        if not st['in_cases']:
            return {'t': 'h', 'level': 3, 'kr': label, 'en': '', 'kind': 'subpoint'}
    return {'t': 'p', 'text': t}

CLS = {'heading': cls_heading, 'named': cls_named, 'font': cls_font}

def detect_mode(d):
    styles = Counter(p.style.name for p in d.paragraphs)
    if any('Heading' in s for s in styles): return 'heading'
    if any(s in styles for s in NAMED): return 'named'
    return 'font'

def save_image(doc, rid, order, seq, placed_parts):
    part = doc.part.related_parts.get(rid)
    if part is None: return None
    placed_parts.add(part.partname)
    ext = str(part.partname).rsplit('.', 1)[-1].lower()
    if ext == 'jpg': ext = 'jpeg'
    blob = part.blob
    outdir = os.path.join(IMG_DIR, str(order))
    os.makedirs(outdir, exist_ok=True)
    fn = f'{seq:02d}.{("jpg" if ext=="jpeg" else ext)}'
    path = os.path.join(outdir, fn)
    try:
        im = Image.open(io.BytesIO(blob))
        if im.mode in ('P', 'RGBA') and ext == 'jpeg': im = im.convert('RGB')
        if im.width > 1500:
            im = im.resize((1500, round(im.height * 1500 / im.width)))
        im.save(path)
    except Exception:
        with open(path, 'wb') as f: f.write(blob)
    return f'img/{order}/{fn}'

def el_image_rids(el):
    rids = []
    for x in el.iter():
        if x.tag == qn('a:blip'):
            r = x.get(qn('r:embed')) or x.get(qn('r:link'))
            if r: rids.append(r)
        elif x.tag == V_IMAGEDATA:
            r = x.get(R_ID)
            if r: rids.append(r)
    return rids

def iter_blocks(parent_el):
    """Yield ('p'|'tbl', element) for top-level block items, descending into
    content controls (w:sdt) so nothing is skipped."""
    for child in parent_el.iterchildren():
        tag = child.tag
        if tag == qn('w:p'): yield ('p', child)
        elif tag == qn('w:tbl'): yield ('tbl', child)
        elif tag == qn('w:sdt'):
            cont = child.find(qn('w:sdtContent'))
            if cont is not None: yield from iter_blocks(cont)

def walk(d, mode, order, label):
    """Document-order pass over top-level blocks. Paragraphs → text + inline
    images; tables with images → flattened figure(image+caption); data tables →
    a structured 'table' block so they render as real tables (not fragmented)."""
    content = []; st = {'in_cases': False}; placed = set(); seqbox = [0]
    clsfn = CLS[mode]
    def emit_image(rid):
        seqbox[0] += 1
        src = save_image(d, rid, order, seqbox[0], placed)
        if src: content.append({'t': 'img', 'src': src})
    def emit_para(p_el):
        para = Paragraph(p_el, d)
        t = norm(para.text)
        if t:
            blk = clsfn(para, t, st)
            if blk:
                if blk['t'] == 'p' and CAPTION_RE.match(t): blk = {'t': 'cap', 'text': t}
                content.append(blk)
        for rid in el_image_rids(p_el):
            emit_image(rid)

    for kind, el in iter_blocks(d.element.body):
        if kind == 'p':
            emit_para(el)
        else:  # table
            tbl = Table(el, d)
            has_img = bool(el_image_rids(el))
            try:
                ncols = len(tbl.columns); nrows = len(tbl.rows)
            except Exception:
                ncols = nrows = 1
            if has_img:                       # figure-layout table → flatten cells
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            emit_para(p._p)
            elif nrows >= 2 or ncols >= 2:     # real data table → table block
                rows = []
                for row in tbl.rows:
                    seen = []; cells = []
                    for c in row.cells:
                        if id(c._tc) in seen: continue   # skip merged repeats
                        seen.append(id(c._tc)); cells.append(norm(c.text))
                    rows.append(cells)
                # drop fully-empty rows
                rows = [r for r in rows if any(x for x in r)]
                if rows: content.append({'t': 'table', 'rows': rows})
            else:                              # 1×1 wrapper → flatten paragraphs
                for p in tbl._cells[0].paragraphs if tbl._cells else []:
                    emit_para(p._p)

    # fallback: images embedded in the package but never placed in the body
    extra = [part for part in d.part.package.iter_parts()
             if 'image' in part.content_type and part.partname not in placed]
    if extra:
        content.append({'t': 'h', 'level': 1, 'kr': '그림·사진', 'en': 'Figures', 'kind': 'figs'})
        content.append({'t': 'note', 'text': '※ 아래 그림은 원본 문서에 본문 위치정보 없이 포함된 이미지입니다(서식적용본). '
                        '15장·16장은 동일한 이미지 세트를 공유하므로, 장별 사용 그림 확인이 필요합니다.'})
        n = 0
        for part in sorted(extra, key=lambda x: str(x.partname)):
            n += 1; seqbox[0] += 1
            ext = str(part.partname).rsplit('.', 1)[-1].lower(); ext = 'jpg' if ext == 'jpeg' else ext
            outdir = os.path.join(IMG_DIR, str(order)); os.makedirs(outdir, exist_ok=True)
            fn = f'{seqbox[0]:02d}.{ext}'; path = os.path.join(outdir, fn)
            try:
                im = Image.open(io.BytesIO(part.blob))
                if im.width > 1500: im = im.resize((1500, round(im.height*1500/im.width)))
                im.save(path)
            except Exception:
                with open(path, 'wb') as f: f.write(part.blob)
            content.append({'t': 'cap', 'text': f'[그림 {label}-{n:02d}] (원본 미배치 · 위치 확인 필요)'})
            content.append({'t': 'img', 'src': f'img/{order}/{fn}'})
    hl = [c['level'] for c in content if c['t'] == 'h']
    if hl and mode == 'heading':
        shift = min(hl) - 1
        for c in content:
            if c['t'] == 'h' and c['kind'] != 'figs': c['level'] -= shift
    return content

# ---- case keywords ----
PLACES = ['청계천','청계고가','서울로','마포','구로','구로공단','G밸리','강남','여의도','영동','상계','목동','잠실',
          '난지도','수도권매립지','한강','DMC','마곡','성수','문래','북한산','남산','광화문','을지로','종로','세운']
KW_OVERRIDE = [
    ('한강종합개발사업과 상수원', ['한강종합개발', '상수원 관리']),
    ('유수율 제고사업', ['유수율 제고사업', '과학적 관망관리']),
    ('상수도 GIS', ['상수도 GIS', '유틸리티 모델']),
    ('쓰레기 수수료 종량제', ['쓰레기 종량제']),
    ('자원회수시설 공동소각', ['자원회수시설', '공동소각']),
    ('수도권 매립지 조성', ['수도권 매립지', '광역 폐기물관리']),
    ('원전 하나 줄이기', ['원전 하나 줄이기']),
    ('건물 에너지 효율화', ['건물 에너지 효율화']),
    ('한강종합개발사업 (1982', ['한강종합개발', '1982~1986']),
    ('하수도 분류식화', ['하수도 분류식화', '빗물 저류']),
    ('한강 수계 수자원 통합', ['한강수계 통합관리', '가뭄 대응']),
    ('점·선의 비움', ['점·선 녹지', '녹지 네트워크']),
    ('강을 되살리다', ['하천 복원']),
]
JOSA = re.compile(r'(으로|에서|에게|로서|로써|를|을|로|와|과|의|이|가|은|는|에)$')
VERB_END = re.compile(r'(다|라|자|었다|ㄴ다|되다|하다|짓다|들다|이다)$')
def _ct(tok):
    tok = tok.strip(' .,’‘“”\'"·-—()[]'); tok = JOSA.sub('', tok); return tok.strip()
def case_keywords(kr):
    for sub, kw in KW_OVERRIDE:
        if sub in kr: return kw
    kws = []
    for pr in re.findall(r'[((]\s*([^))]+?)\s*[))]', kr):
        if pr.strip() and pr.strip() not in kws: kws.append(pr.strip())
    for ac in re.findall(r'[A-Za-z][A-Za-z\-]{1,}', kr):
        if ac.lower() in ('point','core','and') or len(ac) < 2: continue
        if ac not in kws: kws.append(ac)
    for yr in re.findall(r'\d{4}(?:\s*[~–-]\s*\d{4})?', kr):
        if yr not in kws: kws.append(yr)
    for pl in PLACES:
        if pl in kr and pl not in kws: kws.append(pl)
    body = re.sub(r'^(?:핵심\s*(?:정책)?\s*[①-⑩\d]|연결\s*\d|Core[^:：]*\d)\s*[:：]?\s*', '', kr).strip()
    body = re.sub(r'[((].*?[))]', '', body)
    parts = re.split(r'[:：]', body)
    phrase = parts[-1] if len(parts) > 1 else body
    for tok in re.split(r'\s*[·,/→]\s*|\s*그리고\s*|(?:와|과|및)(?=\s)', phrase):
        tok = _ct(tok)
        if not tok or len(tok) > 14 or VERB_END.search(tok) or tok.count(' ') >= 2: continue
        if tok not in kws: kws.append(tok)
    if not kws and len(parts) > 1:
        lead = _ct(parts[0].split()[0]) if parts[0].split() else ''
        if lead: kws.append(lead)
    out = [k for k in kws if not any(k != o and k in o for o in kws)]
    return out[:4] or [_ct(phrase)[:16] or kr[:16]]

def main():
    files = glob.glob(SRC + '/*/*.docx')
    bykey = {norm(os.path.basename(f)): f for f in files}
    chapters = []
    for idx, (key, part, partname, titleKR, titleEN, label) in enumerate(META):
        match = next((f for k, f in bykey.items() if key in k), None)
        if not match: print('MISSING', key); continue
        d = docx.Document(match)
        mode = detect_mode(d)
        content = walk(d, mode, idx, label)
        headings = [c for c in content if c['t'] == 'h']
        cases = [{'kr': c['kr'], 'en': c['en'], 'keywords': case_keywords(c['kr'])}
                 for c in headings if c['kind'] == 'core']
        nimg = sum(1 for c in content if c['t'] == 'img')
        ntbl = sum(1 for c in content if c['t'] == 'table')
        chars = sum(len(c['text']) for c in content if c['t'] == 'p')
        chapters.append({
            'order': idx, 'label': label, 'num': key, 'part': part, 'partname': partname,
            'titleKR': titleKR, 'titleEN': titleEN, 'mode': mode,
            'file': norm(os.path.basename(match)), 'content': content,
            'headings': [{'level': h['level'], 'kr': h['kr'], 'en': h['en'], 'kind': h['kind']} for h in headings],
            'cases': cases, 'caseCount': len(cases), 'chars': chars, 'images': nimg, 'tables': ntbl,
        })
        print(f"{label:5s} {titleKR:22s} {mode:7s} blk={len(content):4d} cases={len(cases)} imgs={nimg} tbls={ntbl} chars={chars}")
    out = {'chapters': chapters,
           'meta': {'publisher': '서울연구원 글로벌 연구협력센터',
                    'titleKR': '성장하는 도시를 위한 도시계획',
                    'titleEN': 'Urban Planning for Growing Cities'}}
    with open(os.path.join(HERE, 'data.json'), 'w', encoding='utf-8') as fp:
        json.dump(out, fp, ensure_ascii=False)
    print('TOTAL', len(chapters), 'chapters,', sum(c['caseCount'] for c in chapters), 'cases,',
          sum(c['images'] for c in chapters), 'images')

if __name__ == '__main__':
    main()
